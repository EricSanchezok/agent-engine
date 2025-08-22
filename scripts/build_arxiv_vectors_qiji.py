import re
import docx
from pathlib import Path

def find_arxiv_ids_in_docx(file_path):
    try:
        document = docx.Document(file_path)
        arxiv_ids = []
        arxiv_pattern = re.compile(r'https?://arxiv\.org/[\w/.-?=&]+')

        for para in document.paragraphs:
            found_ids = arxiv_pattern.findall(para.text)
            arxiv_ids.extend(found_ids)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        found_ids = arxiv_pattern.findall(para.text)
                        arxiv_ids.extend(found_ids)

        rels = document.part.rels
        for rel in rels:
            if rels[rel].reltype.endswith('hyperlink'):
                link_url = rels[rel]._target
                if arxiv_pattern.match(link_url):
                    arxiv_ids.append(link_url)
                    
        return list(set(arxiv_ids))
    except Exception as e:
        print(f"处理文件 {file_path} 时出错: {e}")
        return []

file_path = Path("database/signals_qiji")
file_names = [f"{file_path}/{f.name}" for f in file_path.iterdir() if f.is_file()]
print(f"file_names: {len(file_names)}")

all_found_ids = []
for file in file_names:
    all_found_ids.extend(find_arxiv_ids_in_docx(file))

unique_ids = list(set(all_found_ids))
print(f"找到的独立链接总数: {len(unique_ids)}")
# print(unique_ids)


def extract_arxiv_ids(urls):
    id_pattern = re.compile(r"(\d+\.\d+)")
    
    found_ids = []
    for url in urls:
        match = id_pattern.search(url)
        if match:
            arxiv_id = match.group(1)
            found_ids.append(arxiv_id)
            
    # 使用 set 去除重复的 ID，然后转换回 list
    unique_ids = list(set(found_ids))
    
    return unique_ids

final_ids = extract_arxiv_ids(unique_ids)
print(f"成功提取并去重后，得到 {len(final_ids)} 个唯一ID：")
# print(final_ids)


import json

output_filename = Path("database/arxiv_qiji/arxiv_qiji_ids.json")
output_filename.parent.mkdir(parents=True, exist_ok=True)

with open(output_filename, "w", encoding="utf-8") as f:
    json.dump(final_ids, f, ensure_ascii=False, indent=4)

print(f"output_filename: {output_filename}")
