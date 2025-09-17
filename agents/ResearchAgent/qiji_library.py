import sys
import os
from pathlib import Path

# Add project root to Python path
current_file = Path(__file__).resolve()
project_root = current_file
while project_root.parent != project_root:
    if (project_root / "pyproject.toml").exists():
        break
    project_root = project_root.parent
sys.path.insert(0, str(project_root))


from dotenv import load_dotenv
import os
from pprint import pprint
import asyncio
import re
import docx
import numpy as np
from pathlib import Path
from typing import List, Dict, Set, Optional, Union
from datetime import datetime
import pyinstrument

# Agent Engine imports
from agent_engine.llm_client import QzClient
from agent_engine.agent_logger import AgentLogger
from agent_engine.memory.e_memory import EMemory, Record

# Core imports
from core.arxiv_fetcher import ArxivFetcher, ArxivPaper

# Local imports
from agents.ResearchAgent.config import QIJI_LIBRARY_DIR, PDF_STROAGE_DIR

load_dotenv()

logger = AgentLogger(__name__)

USE_ERIC_VPN = os.getenv('USE_ERIC_VPN', 'false').lower() == 'true'

def find_arxiv_ids_in_docx(file_path: str) -> List[str]:
    """Extract arXiv IDs from a docx file."""
    try:
        document = docx.Document(file_path)
        arxiv_ids = []
        arxiv_pattern = re.compile(r'https?://arxiv\.org/[\w/.-?=&]+')

        for para in document.paragraphs:
            found_ids = arxiv_pattern.findall(para.text)
            arxiv_ids.extend(found_ids)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        found_ids = arxiv_pattern.findall(para.text)
                        arxiv_ids.extend(found_ids)

        rels = document.part.rels
        for rel in rels:
            if rels[rel].reltype.endswith('hyperlink'):
                link_url = rels[rel]._target
                if arxiv_pattern.match(link_url):
                    arxiv_ids.append(link_url)
                    
        return list(set(arxiv_ids))
    except Exception as e:
        logger.error(f"Error processing {file_path}: {e}")
        return []

def extract_arxiv_ids(urls: List[str]) -> List[str]:
    """Extract arXiv IDs from URLs."""
    id_pattern = re.compile(r"(\d+\.\d+)")
    
    found_ids = []
    for url in urls:
        match = id_pattern.search(url)
        if match:
            arxiv_id = match.group(1)
            found_ids.append(arxiv_id)
            
    unique_ids = list(set(found_ids))
    return unique_ids

class QijiLibrary:
    def __init__(self):
        self.embedding_client = QzClient(api_key=os.getenv('QZ_API_KEY'), base_url=os.getenv('QWEN3_EMBEDDING_8B_H100_URL') if not USE_ERIC_VPN else os.getenv('ERIC_VPN_URL') + os.getenv('QWEN3_EMBEDDING_8B_H100_PROXY_ROUTE'))
        self.reranker_client = QzClient(api_key=os.getenv('QZ_API_KEY'), base_url=os.getenv('QWEN3_RERANKER_8B_H100_URL') if not USE_ERIC_VPN else os.getenv('ERIC_VPN_URL') + os.getenv('QWEN3_RERANKER_8B_H100_PROXY_ROUTE'))
        self.embedding_model = "eric-qwen3-embedding-8b"
        self.reranker_model = "eric-qwen3-reranker-8b"
        self.qiji_memory = EMemory(name='qiji_memory', persist_dir=QIJI_LIBRARY_DIR)
        self.arxiv_fetcher = ArxivFetcher(pdf_storage_dir=PDF_STROAGE_DIR)
        
    async def update_memory(self, qiji_articles_dir: str = None, max_concurrent: int = 1) -> Dict[str, int]:
        """
        Update qiji memory by processing all docx files in qiji_articles directory.
        
        Args:
            qiji_articles_dir: Directory containing qiji articles (defaults to database/qiji_articles)
            max_concurrent: Maximum concurrent embedding tasks
            
        Returns:
            Dictionary with statistics about the update process
        """
        if qiji_articles_dir is None:
            qiji_articles_dir = QIJI_LIBRARY_DIR
        
        qiji_articles_path = Path(qiji_articles_dir)
        if not qiji_articles_path.exists():
            logger.error(f"Qiji articles directory not found: {qiji_articles_path}")
            return {"error": "Directory not found"}
        
        # Get all docx files
        docx_files = list(qiji_articles_path.glob("*.docx"))
        if not docx_files:
            logger.warning(f"No docx files found in {qiji_articles_path}")
            return {"processed_files": 0, "total_arxiv_ids": 0, "new_papers": 0, "skipped_papers": 0}
        
        logger.info(f"Found {len(docx_files)} docx files to process")
        
        # Extract all arXiv IDs from docx files
        all_arxiv_ids = []
        for docx_file in docx_files:
            logger.info(f"Processing file: {docx_file.name}")
            found_urls = find_arxiv_ids_in_docx(str(docx_file))
            extracted_ids = extract_arxiv_ids(found_urls)
            all_arxiv_ids.extend(extracted_ids)
            logger.info(f"Found {len(extracted_ids)} arXiv IDs in {docx_file.name}")
        
        # Remove duplicates
        unique_arxiv_ids = list(set(all_arxiv_ids))
        logger.info(f"Total unique arXiv IDs found: {len(unique_arxiv_ids)}")
        
        if not unique_arxiv_ids:
            logger.warning("No arXiv IDs found in any docx files")
            return {"processed_files": len(docx_files), "total_arxiv_ids": 0, "new_papers": 0, "skipped_papers": 0}
        
        # Check which papers already exist in memory and have vectors
        existing_papers = set()
        for arxiv_id in unique_arxiv_ids:
            if self.qiji_memory.exists(arxiv_id) and self.qiji_memory.has_vector(arxiv_id):
                existing_papers.add(arxiv_id)
        
        new_arxiv_ids = [arxiv_id for arxiv_id in unique_arxiv_ids if arxiv_id not in existing_papers]
        logger.info(f"Found {len(existing_papers)} existing papers with vectors, {len(new_arxiv_ids)} new papers to process")
        
        if not new_arxiv_ids:
            logger.info("All papers already exist in memory")
            return {
                "processed_files": len(docx_files),
                "total_arxiv_ids": len(unique_arxiv_ids),
                "new_papers": 0,
                "skipped_papers": len(existing_papers)
            }
        
        # Fetch new papers using ArxivFetcher
        logger.info(f"Fetching {len(new_arxiv_ids)} new papers from arXiv")
        try:
            papers = await self.arxiv_fetcher.search_papers(paper_ids=new_arxiv_ids)
            logger.info(f"Successfully fetched {len(papers)} papers")
        except Exception as e:
            logger.error(f"Error fetching papers: {e}")
            return {"error": f"Failed to fetch papers: {e}"}
        
        if not papers:
            logger.warning("No papers were fetched")
            return {
                "processed_files": len(docx_files),
                "total_arxiv_ids": len(unique_arxiv_ids),
                "new_papers": 0,
                "skipped_papers": len(existing_papers)
            }
        
        # Create semaphore for concurrent embedding
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def embed_and_store_paper(paper):
            """Embed and store a single paper."""
            async with semaphore:
                try:
                    # Check if paper already exists and has vector (double-check for race conditions)
                    if self.qiji_memory.exists(paper.id) and self.qiji_memory.has_vector(paper.id):
                        logger.debug(f"Paper {paper.id} already exists with vector, skipping")
                        return False
                    
                    # Create embedding for paper summary
                    summary = paper.summary
                    if not summary:
                        logger.warning(f"No summary available for paper {paper.id}")
                        return False
                    
                    logger.debug(f"Creating embedding for paper {paper.id}")
                    embedding = await self.embedding_client.embedding(
                        text=summary,
                        model_name=self.embedding_model
                    )
                    
                    if not embedding:
                        logger.error(f"Failed to create embedding for paper {paper.id}")
                        return False
                    
                    vector = embedding
                    
                    # Create record
                    record = Record(
                        id=paper.id,
                        vector=vector,
                        content=summary
                    )
                    
                    # Store in memory
                    success = self.qiji_memory.add(record)
                    if success:
                        logger.debug(f"Successfully stored paper {paper.id}")
                        return True
                    else:
                        logger.error(f"Failed to store paper {paper.id}")
                        return False
                        
                except Exception as e:
                    logger.error(f"Error processing paper {paper.id}: {e}")
                    return False
        
        # Process papers concurrently
        logger.info(f"Starting concurrent embedding and storage (max_concurrent={max_concurrent})")
        tasks = [embed_and_store_paper(paper) for paper in papers]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Count results
        successful_count = 0
        failed_count = 0
        for result in results:
            if isinstance(result, Exception):
                failed_count += 1
                logger.error(f"Task failed with exception: {result}")
            elif result:
                successful_count += 1
            else:
                failed_count += 1
        
        logger.info(f"Update completed: {successful_count} papers added, {failed_count} papers failed")
        
        return {
            "processed_files": len(docx_files),
            "total_arxiv_ids": len(unique_arxiv_ids),
            "new_papers": successful_count,
            "skipped_papers": len(existing_papers),
            "failed_papers": failed_count
        }
    
    async def find_minimum_distance(self, input_data: Union[ArxivPaper, List[float]]) -> Optional[float]:
        """
        Find the minimum distance between input vector and all vectors in qiji_memory.
        
        Args:
            input_data: Either an ArxivPaper object or a vector (list of floats)
            
        Returns:
            Minimum distance if found, None if no vectors exist in memory
        """
        try:
            # Get input vector
            if isinstance(input_data, ArxivPaper):
                # Create embedding for paper summary
                if not input_data.summary:
                    logger.warning(f"No summary available for paper {input_data.id}")
                    return None
                
                logger.debug(f"Creating embedding for paper {input_data.id}")
                embedding = await self.embedding_client.embedding(
                    text=input_data.summary,
                    model_name=self.embedding_model
                )
                
                if not embedding:
                    logger.error(f"Failed to create embedding for paper {input_data.id}")
                    return None
                
                input_vector = embedding
            else:
                # Assume input_data is a vector
                input_vector = input_data
            
            # Use EMemory's search_similar_records to find the most similar record
            # Since ChromaDB returns results sorted by distance, we only need k=1
            similar_records = self.qiji_memory.search_similar_records(
                query_vector=input_vector,
                k=1
            )
            
            if not similar_records:
                logger.warning("No similar records found in qiji_memory")
                return None
            
            # The first (and only) result has the minimum distance
            _, min_distance = similar_records[0]
            
            logger.debug(f"Minimum distance found: {min_distance:.4f}")
            return float(min_distance)
            
        except Exception as e:
            logger.error(f"Error in find_minimum_distance: {e}")
            return None
    
    async def find_minimum_distances_batch(self, input_data_list: Union[List[ArxivPaper], List[List[float]]]) -> List[Optional[float]]:
        """
        Find minimum distances for a batch of inputs.
        
        Args:
            input_data_list: List of ArxivPaper objects or vectors (list of floats)
            
        Returns:
            List of minimum distances, None for failed computations
        """
        async def _process_single_input(input_data: Union[ArxivPaper, List[float]], index: int) -> Optional[float]:
            """Process a single input and return its minimum distance."""
            try:
                # Get input vector
                if isinstance(input_data, ArxivPaper):
                    # Create embedding for paper summary
                    if not input_data.summary:
                        logger.warning(f"No summary available for paper {input_data.id} (index {index})")
                        return None
                    
                    logger.debug(f"Creating embedding for paper {input_data.id} (index {index})")
                    embedding = await self.embedding_client.embedding(
                        text=input_data.summary,
                        model_name=self.embedding_model
                    )
                    
                    if not embedding:
                        logger.error(f"Failed to create embedding for paper {input_data.id} (index {index})")
                        return None
                    
                    input_vector = embedding
                else:
                    # Assume input_data is a vector
                    if input_data is None:
                        logger.warning(f"Input vector is None for index {index}")
                        return None
                    input_vector = input_data
                
                # Use EMemory's search_similar_records to find the most similar record
                similar_records = self.qiji_memory.search_similar_records(
                    query_vector=input_vector,
                    k=1
                )
                
                if not similar_records:
                    logger.warning(f"No similar records found for input {index}")
                    return None
                
                # The first (and only) result has the minimum distance
                _, min_distance = similar_records[0]
                
                logger.debug(f"Input {index}: minimum distance = {min_distance:.4f}")
                return float(min_distance)
                
            except Exception as e:
                logger.error(f"Error processing input {index}: {e}")
                return None
                
        try:
            if not input_data_list:
                logger.warning("Empty input list provided")
                return []
            
            logger.info(f"Processing batch of {len(input_data_list)} inputs")
            
            # Process inputs concurrently
            tasks = []
            for i, input_data in enumerate(input_data_list):
                task = _process_single_input(input_data, i)
                tasks.append(task)
            
            # Execute all tasks concurrently
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Process results
            distances = []
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    logger.error(f"Error processing input {i}: {result}")
                    distances.append(None)
                else:
                    distances.append(result)
            
            successful_count = sum(1 for d in distances if d is not None)
            logger.info(f"Batch processing completed: {successful_count}/{len(input_data_list)} successful")
            
            return distances
            
        except Exception as e:
            logger.error(f"Error in find_minimum_distances_batch: {e}")
            return [None] * len(input_data_list)
    



async def update_qiji_memory():
    qiji_library = QijiLibrary()
    await qiji_library.update_memory()

async def test():
    from agents.ResearchAgent.arxiv_database import ArxivDatabase
    arxiv_database = ArxivDatabase()
    qiji_library = QijiLibrary()
    vectors = arxiv_database.get_vectors_by_date(datetime(2025, 9, 11))
    distance = await qiji_library.find_minimum_distances_batch(vectors[:10])
    print(len(distance))
    print(distance)

if __name__ == "__main__":
    asyncio.run(update_qiji_memory())