import os
import re
import docx
from pathlib import Path
import json
from typing import List, Tuple
import asyncio
from dotenv import load_dotenv
from uuid import uuid4
from sklearn.metrics.pairwise import cosine_similarity
from copy import deepcopy
import numpy as np
import datetime
import pytz

# A2A framework imports
from a2a.server.agent_execution import AgentExecutor, RequestContext
from a2a.server.events import EventQueue
from a2a.server.tasks import InMemoryTaskStore
from a2a.server.request_handlers import DefaultRequestHandler
from a2a.server.apps import A2AStarletteApplication
from a2a.types import AgentCard
from a2a.utils import (
    new_agent_text_message, new_agent_parts_message, new_artifact,
    new_data_artifact, new_task, new_text_artifact
)
from a2a.types import (
    Artifact, Message, Role, Task, TaskStatus, TaskState, 
    FilePart, Part, FileWithBytes, TextPart, MessageSendParams
)

# AgentEngine imports
from agent_engine.agent import BaseA2AAgent
from agent_engine.llm_client import AzureClient
from agent_engine.agent import SkillIdentifier
from agent_engine.prompt import PromptLoader
from agent_engine.agent_logger import AgentLogger
from agent_engine.memory import Memory
from agent_engine.utils import get_relative_path_from_current_file, get_current_file_dir

# Core imports
from core.arxiv import ArXivFetcher, Paper, CATEGORIES_QUERY_STRING
from core.utils import DateFormatter

# Local imports
from agents.PaperFilterAgent.config import AGENT_CARD

logger = AgentLogger(__name__)

load_dotenv()

class PaperFilterAgent(BaseA2AAgent):
    def __init__(self):
        super().__init__(agent_card=AGENT_CARD)
        self.llm_client = AzureClient(api_key=os.getenv('AZURE_API_KEY'))
        self.skill_identifier = SkillIdentifier(llm_client=self.llm_client, model_name='o3-mini')
        self.prompt_loader = PromptLoader(get_relative_path_from_current_file('prompts.yaml'))
        self.arxiv_memory = Memory(name='arxiv_memory', db_path=get_current_file_dir() / 'database' / 'arxiv_memory.db')
        self.arxiv_qiji_memory = Memory(name='arxiv_qiji_memory', db_path=get_current_file_dir() / 'database' / 'arxiv_qiji_memory.db')
        self.arxiv_fetcher = ArXivFetcher()
        self.date_formatter = DateFormatter()
        self.semaphore = asyncio.Semaphore(32)

    async def execute(self, context: RequestContext, event_queue: EventQueue) -> None:
        task_id = context.task_id if context.task_id else str(uuid4())
        context_id = context.context_id if context.context_id else str(uuid4())
        user_input = context.get_user_input()

        skill_id, reason = await self.skill_identifier.invoke(user_input, AGENT_CARD.skills)
        logger.info(f"Skill ID: {skill_id}, Reason: {reason}")

        if skill_id == 'filter_and_recommend':
            try:
                arxiv_json = json.loads(user_input)
                arxiv_ids = arxiv_json.get('arxiv_ids', [])
                max_recommendations = arxiv_json.get('max_recommendations', 16)
            except Exception as e:
                await self._task_failed(context, event_queue, f"Can't parse the input: {e}")
                return

            if not arxiv_ids:
                await self._task_failed(context, event_queue, f"No arxiv ids provided")
                return

            try:
                papers: List[Paper] = await self.arxiv_fetcher.search(id_list=arxiv_ids)
            except Exception as e:
                await self._task_failed(context, event_queue, f"Can't fetch the papers: {e}")
                return
            
            if not papers:
                await self._task_failed(context, event_queue, f"No papers found")
                return
            
            tasks = [self._embed_paper(idx, paper) for idx, paper in enumerate(papers)]
            result = await asyncio.gather(*tasks, return_exceptions=True)
            
            valid_papers: List[Paper] = []
            valid_vectors: List[List[float]] = []
            for item in result:
                # 跳过异常
                if isinstance(item, Exception):
                    logger.error(f"Embedding task returned exception: {item}")
                    continue
                paper, vector = item  # type: ignore[misc]
                if vector is not None:
                    valid_papers.append(paper)
                    valid_vectors.append(vector)
            
            if not valid_papers:
                await self._task_failed(context, event_queue, f"No valid papers found")
                return

            qiji_vectors = list(self.arxiv_qiji_memory.get_all_vectors().values())

            vectors_np = np.array(valid_vectors, dtype=np.float32)
            qiji_vectors_np = np.array(qiji_vectors, dtype=np.float32)

            similarity_matrix = cosine_similarity(vectors_np, qiji_vectors_np)
            max_similarities = np.max(similarity_matrix, axis=1)

            papers_with_similarity: List[Tuple[Paper, float]] = [
                (paper, similarity) for paper, similarity in zip(valid_papers, max_similarities)
            ]

            sorted_papers = sorted(papers_with_similarity, key=lambda x: x[1], reverse=True)

            top_papers = sorted_papers[:max_recommendations]

            result_ids = [paper.id.replace('_', '.') for paper, _ in top_papers]

            json_content = json.dumps(result_ids, ensure_ascii=False, indent=4)
            parts = [Part(root=TextPart(text=json_content))]

            artifact = Artifact(
                artifact_id=str(uuid4()),
                parts=parts,
                name="arxiv_filter_papers",
                description=f"ArXiv筛选的论文，共{len(top_papers)}篇",
            )
            
            message = Message(
                role=Role.agent,
                task_id=task_id,
                message_id=str(uuid4()),
                content_id=context_id,
                parts=parts
            )

            task = Task(
                id=task_id,
                contextId=context_id,
                artifacts=[artifact],
                history=[message],
                status=TaskStatus(
                    state=TaskState.completed,
                    timestamp=datetime.datetime.now(pytz.timezone('Asia/Shanghai')).replace(microsecond=0).isoformat(),
                ),
            )
            await self._put_event(event_queue, task)
            logger.info(f"Collection results sent - {len(top_papers)} papers metadata returned")

        else:
            await self._task_failed(context, event_queue, f"Can't find the skill: {skill_id}")
            return

    async def _embed_paper(self, idx, paper: Paper) -> Tuple[Paper, List[float]]:
        async with self.semaphore:
            logger.debug(f"Processing paper {idx} with {paper.id}")
            summary = paper.info.get('summary', '')
            try:
                vector, _ = self.arxiv_memory.get_by_content(summary)
                if vector is None:
                    vector = await self.llm_client.embedding(summary, model_name='text-embedding-3-large')
                    self.arxiv_memory.add(summary, vector)
                return paper, vector
            except Exception as e:
                logger.error(f"Error embedding paper {idx}: {e}")
                return paper, None

def find_arxiv_ids_in_docx(file_path):
    try:
        document = docx.Document(file_path)
        arxiv_ids = []
        arxiv_pattern = re.compile(r'https?://arxiv\.org/[\w/.-?=&]+')

        for para in document.paragraphs:
            found_ids = arxiv_pattern.findall(para.text)
            arxiv_ids.extend(found_ids)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        found_ids = arxiv_pattern.findall(para.text)
                        arxiv_ids.extend(found_ids)

        rels = document.part.rels
        for rel in rels:
            if rels[rel].reltype.endswith('hyperlink'):
                link_url = rels[rel]._target
                if arxiv_pattern.match(link_url):
                    arxiv_ids.append(link_url)
                    
        return list(set(arxiv_ids))
    except Exception as e:
        print(f"Process {file_path} error: {e}")
        return []

def extract_arxiv_ids(urls):
    id_pattern = re.compile(r"(\d+\.\d+)")
    
    found_ids = []
    for url in urls:
        match = id_pattern.search(url)
        if match:
            arxiv_id = match.group(1)
            found_ids.append(arxiv_id)
            
    unique_ids = list(set(found_ids))
    
    return unique_ids

async def build_arxiv_qiji_memory():
    agent = PaperFilterAgent()

    file_path = Path("database/signals_qiji")
    file_names = [f"{file_path}/{f.name}" for f in file_path.iterdir() if f.is_file()]
    logger.info(f"file_names: {len(file_names)}")

    all_found_ids = []
    for file in file_names:
        all_found_ids.extend(find_arxiv_ids_in_docx(file))

    unique_ids = list(set(all_found_ids))

    final_ids = extract_arxiv_ids(unique_ids)

    output_filename = Path("database/arxiv_qiji/arxiv_qiji_ids.json")
    output_filename.parent.mkdir(parents=True, exist_ok=True)

    with open(output_filename, "w", encoding="utf-8") as f:
        json.dump(final_ids, f, ensure_ascii=False, indent=4)

    papers: List[Paper] = await agent.arxiv_fetcher.search(id_list=final_ids)

    logger.info(f"papers: {len(papers)}")

    categories = set()
    for paper in papers:
        categories.update(paper.info['categories'])
    logger.info(categories)

    for paper in papers:
        paper_id = paper.id
        summary = paper.info['summary']
        vector, _ = agent.arxiv_memory.get_by_content(summary)
        if vector is None:
            vector = await agent.llm_client.embedding(summary, model_name='text-embedding-3-large')
        agent.arxiv_qiji_memory.add(summary, vector)
        agent.arxiv_memory.add(summary, vector)
        logger.info(f"Saved {paper_id}")


async def main():
    ids = [
        '2508_15697',
        '2508.15692',
        '2508.15678'
    ]
    agent = PaperFilterAgent()
    await agent.run_user_input(json.dumps({
        "arxiv_ids": ids,
        "max_recommendations": 2
    }, ensure_ascii=False, indent=4))


if __name__ == "__main__":
    asyncio.run(build_arxiv_qiji_memory())